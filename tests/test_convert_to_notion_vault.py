import os
import unittest
import tempfile
import shutil
import docx
from scripts.convert_to_notion_vault import docx_table_to_markdown, process_docx

class TestConvertToNotionVault(unittest.TestCase):
    def setUp(self):
        self.test_dir = tempfile.mkdtemp()
        self.output_dir = tempfile.mkdtemp()

    def tearDown(self):
        shutil.rmtree(self.test_dir)
        shutil.rmtree(self.output_dir)

    def test_docx_table_to_markdown(self):
        # Create a document and a table inside it
        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Col A"
        table.cell(0, 1).text = "Col B"
        table.cell(1, 0).text = "Val 1"
        table.cell(1, 1).text = "Val 2"

        md_output = docx_table_to_markdown(table)
        self.assertIn("| Col A | Col B |", md_output)
        self.assertIn("| --- | --- |", md_output)
        self.assertIn("| Val 1 | Val 2 |", md_output)

    def test_process_docx_generates_markdown_and_sub_pages(self):
        # Generate a docx with paragraphs and tables to verify full integration
        doc_path = os.path.join(self.test_dir, "MyDocument.docx")
        doc = docx.Document()
        doc.add_paragraph("Paragraph 1")

        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H1"
        table.cell(0, 1).text = "H2"
        table.cell(1, 0).text = "D1"
        table.cell(1, 1).text = "D2"

        doc.add_paragraph("Paragraph 2")
        doc.save(doc_path)

        # Process the DOCX file
        process_docx(doc_path, self.output_dir)

        # Check index markdown exists
        index_md = os.path.join(self.output_dir, "MyDocument.md")
        self.assertTrue(os.path.exists(index_md))

        # Check sub-page exists
        sub_page_path = os.path.join(self.output_dir, "MyDocument", "MyDocument-Part-1.md")
        self.assertTrue(os.path.exists(sub_page_path))

        # Check sub-page content includes paragraphs and the table in order
        with open(sub_page_path, "r", encoding="utf-8") as f:
            content = f.read()
            self.assertIn("Paragraph 1", content)
            self.assertIn("| H1 | H2 |", content)
            self.assertIn("| --- | --- |", content)
            self.assertIn("| D1 | D2 |", content)
            self.assertIn("Paragraph 2", content)

if __name__ == "__main__":
    unittest.main()
